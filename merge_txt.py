from docx import Document
import os
import re


path = './'
direct = os.listdir(path)

from docx.shared import Pt

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

numbers = re.compile(r'(\d+)')
def numericalSort(value):
    parts = numbers.split(value)
    parts[1::2] = map(int, parts[1::2])
    return parts

for i in sorted(direct, key=numericalSort):
    try:            
        myfile = open('./'+i).read()
        i = i.replace('.txt','')
        document.add_heading(i, 0)
        p = document.add_paragraph(myfile)
        document.add_page_break()
        # document.save('./docx/'+ i +'.docx')
    except Exception as e:
        print(e)
document.save('./final/result.docx')